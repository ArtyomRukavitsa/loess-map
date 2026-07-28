# -*- coding: utf-8 -*-
# Генерирует единый Word-документ поставки (Обзор + Методология + Словарь колонок).
# Стиль — академический (Times New Roman), таблицы с рамками и заголовочной заливкой.
# Выход: delivery/Документация_БД_лёссовые_разрезы.docx
import os, sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__)); DELIV = os.path.join(HERE, "delivery")
NAVY = RGBColor(0x1F, 0x38, 0x64); GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
# базовый стиль — Times New Roman 11
st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
for m in (doc.sections[0].left_margin, ):
    pass
sec = doc.sections[0]; sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), hexcolor); tcPr.append(sh)

def heading(text, level, color=NAVY):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color; r.font.name = "Times New Roman"
    return h

def para(text="", bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        if color: r.font.color.rgb = color
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); r = p.add_run(text); r.font.size = Pt(11); return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        rp = c.paragraphs[0].add_run(htext); rp.bold = True; rp.font.size = Pt(10)
        rp.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "1F3864")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rp = cells[i].paragraphs[0].add_run(str(val)); rp.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ---------- ТИТУЛ ----------
para("Институт географии РАН · Yandex Cloud", italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("База данных лёссовых и четвертичных разрезов"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
para("Автоматическое извлечение из архива научных публикаций", italic=True, size=13,
     color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Версия 1.0 · 14 июля 2026", align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=14)

# ---------- 1. ОБЗОР ----------
heading("1. Обзор", 1)
para("Структурированная база данных разрезов лёссовых и четвертичных отложений "
     "(местоположение, мощность, стратиграфия, методы датирования), извлечённая автоматически "
     "из архива русскоязычных научных публикаций (~10 000 файлов, ~58 ГБ: PDF-сканы, DjVu и "
     "фотографии страниц книг). Цель — заполнить пробелы в международных базах по разрезам "
     "на территории бывшего СССР.")

heading("Состав поставки", 2)
table(["Файл", "Содержимое"],
      [["loess_sections_full.xlsx", "Все 14 701 разрезов (gold-схема + n_sources + confidence_tier)"],
       ["loess_sections_core.xlsx", "686 разрезов «надёжного ядра» (tier A) — для первичного использования"],
       ["loess_sections_international_bonus.xlsx", "6 147 зарубежных разрезов (бонус, вне ТЗ)"],
       ["Интерактивная карта", "github.com/ArtyomRukavitsa/loess-map (Streamlit)"]],
      widths=[6.5, 9.5])

heading("Ключевые цифры", 2)
table(["Показатель", "Значение"],
      [["Уникальных разрезов", "14 701 (консолидировано из 28 629 сырых записей)"],
       ["С географическими координатами", "2 510"],
       ["Источников (продуктивных документов)", "≈629 (490 PDF/скан + 149 фото-публикаций)"],
       ["Методы датирования", "магнитостратиграфия 3 744 · ¹⁴C 3 432 · OSL 1 925 · TL 1 429 · (U–Th)/He 163"],
       ["Стратиграфия", "верх. плейстоцен 8 744 · сред. 5 657 · ниж. 4 254 · голоцен 4 161"]],
      widths=[6.0, 10.0])

heading("Уровни достоверности (confidence_tier)", 2)
table(["Tier", "Что значит", "Кол-во"],
      [["A", "Локализован (координаты) И подтверждён ≥2 источниками ИЛИ богатый (мощность+датирование+стратиграфия)", "686"],
       ["B", "Локализован, один источник", "1 824"],
       ["C", "Есть данные (отложения+стратиграфия), координаты не привязаны", "10 726"],
       ["D", "Разреженная запись", "1 465"]],
      widths=[1.5, 11.5, 2.0])
para("Рекомендация: начните с loess_sections_core.xlsx (tier A). Для полноты — full "
     "с фильтром по confidence_tier и n_sources.", italic=True, color=GREY)

heading("Как читать данные", 2)
bullet("Пустое / нет данных в источнике → ND (единый маркер по всем полям).")
bullet("n_sources — сколько независимых записей консолидировано в разрез (индикатор надёжности).")
bullet("Координаты (N, E) — по названию ближайшего населённого пункта (уровень города/села, радиус ~5 км).")
bullet("Названия — кириллицей, как в источнике.")

heading("Ограничения (важно)", 2)
bullet("Автоматическое извлечение — возможны ошибки OCR и интерпретации LLM. Данные пригодны "
       "для скрининга; точки для публикации рекомендуется сверять с первоисточником.")
bullet("Геокодировано 17% — остальные привязаны к месту только текстом (обскурные названия, номера разрезов).")
bullet("Мощность (21%) и высота (9%) заполнены редко — так в самих источниках.")
bullet("international_bonus — иностранные объекты, извлечены попутно, не валидированы.")

# ---------- 2. МЕТОДОЛОГИЯ ----------
doc.add_page_break()
heading("2. Методология", 1)
para("Конвейер из независимых стадий. Ниже — сжатое описание.", italic=True, color=GREY)

for title, items in [
    ("2.1. Каталогизация архива", [
        "Рекурсивный обход публичного Яндекс.Диска (~10 000 файлов, 58 ГБ) без скачивания; классификация по типам.",
        "Релевантны: PDF/DjVu-сканы + папки с фотографиями страниц книг (149 публикаций, ~7 500 фото).",
        "Исключены: геологические карты, аэрофотоснимки, атласы-плашки."]),
    ("2.2. OCR (распознавание текста)", [
        "Yandex Vision OCR (асинхронный, модель page-column-sort, языки ru+en).",
        "Даунсэмплинг перед OCR (180 DPI / длинная сторона ≤2200 px) — снятие ограничений на размер (сжатие ~7×).",
        "Фотографии книг обрабатывались потоково в памяти, без хранения исходников."]),
    ("2.3. Извлечение записей (LLM)", [
        "Мульти-модельный роутинг (deepseek-v32, qwen3.6-35b, gpt-oss-120b, qwen3-235b) — обход квот, 0 ошибок 429.",
        "Строгая JSON-схема: на каждое поле обязательная цитата-evidence; нет основания → ND. Модели не выдумывают.",
        "Relevance-gate: проба ~17 страниц; 0 разрезов → документ нерелевантный, остальное не обрабатывается.",
        "Выход: 28 629 сырых записей из ≈629 продуктивных документов."]),
    ("2.4. Постобработка и консолидация", [
        "Фильтр иностранного (латиница + список зарубежных топонимов) → отдельный файл.",
        "Нормализация чисел (диапазоны, «свыше»→«≥», см→м) и локаций (снятие префиксов г./с./оз./р.).",
        "Консолидация по разрезу (локация + административная единица): 28 629 → 14 701 разрез. Число слитых → n_sources."]),
    ("2.5. Геокодинг", [
        "Nominatim (OpenStreetMap) с валидацией: страны СНГ, типы населённых пунктов, bounding-box, подсказка по админ-единице.",
        "Привязано 2 510 разрезов (уровень населённого пункта, радиус ~5 км)."]),
    ("2.6. Маппинг в gold-схему", [
        "Приведение к 32-колоночной схеме заказчика, гигиена пустышек (None/nan/\"\" → ND).",
        "Добавление n_sources и confidence_tier."]),
]:
    heading(title, 2); [bullet(x) for x in items]

heading("Контроль качества", 2)
bullet("evidence-цитата на каждое извлечённое поле (трассируемость к тексту).")
bullet("Relevance-gate против ложных срабатываний на нерелевантных документах.")
bullet("Валидация геокодера против выдуманных/зарубежных координат.")

# ---------- 3. СЛОВАРЬ КОЛОНОК ----------
doc.add_page_break()
heading("3. Словарь колонок", 1)
para("34 колонки: 32 gold-схемы заказчика + 2 служебные. Пустое значение везде — ND.", italic=True, color=GREY)

heading("Gold-схема (колонки заказчика)", 2)
gold = [
    ["1", "ID", "Порядковый идентификатор разреза"],
    ["2–3", "N, E", "Широта/долгота (десятичные градусы), геокодинг по названию места"],
    ["4", "Accuracy (radius, m)", "Радиус точности координат, м (5000 = уровень нас. пункта)"],
    ["5", "Thickness, m", "Мощность отложений, м (диапазон «X–Y», «≥X»)"],
    ["6", "Absolute elevation, m a.s.l.", "Абсолютная высота, м над ур. моря"],
    ["7–9", "Type of excavation; Geomorphological position (modern)", "Тип вскрытия и геоморфопозиция — не извлекались в v1"],
    ["10", "Name of geographic feature", "Название геогр. объекта (= ближайший нас. пункт)"],
    ["11", "Nearest locality", "Ближайший населённый пункт"],
    ["12", "Administrative unit", "Административная единица (область/край/республика)"],
    ["13", "Type of deposits", "loess, loess-like loam, slope deposits, alluvium, till, marine, volcanic ash"],
    ["14", "Stratigraphic position", "Lower / Middle / Upper Pleistocene, Holocene"],
    ["15", "Chronological data available", "Есть ли данные датирования (Yes/No)"],
    ["16", "Dating method", "14C, OSL, TL, magnetostratigraphy, (U–Th)/He"],
    ["17–21", "Number of … dates", "Число дат каждого типа — не подсчитывалось в v1 (счётная задача)"],
    ["22", "Publication 1", "Источник (имя файла публикации в архиве)"],
    ["23–29", "DOI / link; Publication 2–4", "DOI недоступен (архивные сканы до эпохи DOI); доп. источники при консолидации"],
    ["30", "Principal investigator", "Ответственный исследователь — не извлекался в v1"],
    ["31", "Comments", "Число источников консолидации + цитата-evidence"],
    ["32", "Data contributor", "auto-extraction (loess pipeline)"],
]
table(["#", "Колонка", "Описание"], gold, widths=[1.3, 5.2, 9.5])

heading("Служебные колонки", 2)
table(["#", "Колонка", "Описание"],
      [["33", "n_sources", "Сколько независимых записей слито в разрез (индикатор надёжности)"],
       ["34", "confidence_tier", "Уровень достоверности A/B/C/D (см. раздел 1)"]],
      widths=[1.3, 5.2, 9.5])

para()
para("Примечание. Колонки 7–9, 17–21, 23, 30 в версии 1.0 не заполнялись: часть — сознательный "
     "выбор в пользу качества ядрового извлечения (узкая схема точнее), часть — счётные под-задачи "
     "(число дат), часть недоступна в источнике (DOI у публикаций доцифровой эпохи). Могут быть "
     "наполнены в следующей итерации при необходимости.", italic=True, color=GREY)

out = os.path.join(DELIV, "Документация_БД_лёссовые_разрезы.docx")
doc.save(out)
print("Word-документ собран:", os.path.normpath(out))
